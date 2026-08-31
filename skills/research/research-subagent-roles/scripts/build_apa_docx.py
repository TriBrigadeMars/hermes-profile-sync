#!/usr/bin/env python3
"""APA 7 .docx package builder for literature reviews (pre-tested).

Builds one .docx per corpus: Title page -> Abstract -> synthesis body (from
synthesis.md) -> APA 7 References (hanging indents, DOIs) -> Annotated
Bibliography ([type] label + summary for EVERY source).

Usage:
    python build_apa_docx.py <workspace_dir> <output.docx> "<Document Title>"

Workspace layout expected:
    <workspace_dir>/sources.json    - validated corpus (standard record schema)
    <workspace_dir>/synthesis.md    - markdown body; ## = H1, ### = H2, #### = H3
    <workspace_dir>/abstract.md     - optional; auto-generated fallback if absent

Smoke test (run BEFORE real corpora):
    python build_apa_docx.py --smoke-test
"""
import json, os, re, sys, tempfile

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

MONTH_YEAR = "August 2026"


def set_apa_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    doc.core_properties.language = "en-US"


def add_page_number_header(doc):
    """APA: page number flush right in header (OxmlElement — do NOT use makeelement)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for section in doc.sections:
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(instr); run._r.append(f2)
        run.font.name = "Times New Roman"; run.font.size = Pt(12)


def _emit_runs(p, text):
    """Parse **bold**, *italic* markdown-lite into runs."""
    pos = 0
    pat = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    for m in pat.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group(2) is not None:
            r = p.add_run(m.group(2)); r.bold = True
        else:
            r = p.add_run(m.group(3)); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_centered(doc, text, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    _emit_runs(p, text)


def hanging(doc, rich_text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5); pf.first_line_indent = Inches(-0.5)
    _emit_runs(p, rich_text)


def fmt_authors(authors):
    """'Lastname, A. A.' strings -> APA author block. Adds ONE trailing period max."""
    a = [x.strip() for x in (authors or []) if x and x.strip()]
    if not a:
        return ""
    a = [x if x.endswith((".", "?", "!")) else x + "." for x in a]
    if len(a) == 1:
        return a[0]
    if len(a) == 2:
        return a[0] + ", & " + a[1]
    return ", ".join(a[:-1]) + ", & " + a[-1]


def ref_entry(s):
    a = fmt_authors(s.get("authors") or [])
    year = s.get("year") or "n.d."
    title = (s.get("title") or "").strip()
    if title and not title.endswith((".", "?", "!")):
        title += "."
    parts = []
    if a:
        parts.append(a)
    elif s.get("no_author"):
        parts.append(title)  # unsigned work: title moves to author position
        title = ""
    parts.append(f"({year}).")
    if title:
        parts.append(title)
    journal, vol, iss, pg = s.get("journal") or "", str(s.get("volume") or ""), s.get("issue") or "", s.get("pages") or ""
    doi = s.get("doi_or_url") or ""
    stype = (s.get("type") or "").lower()
    if "newspaper" in stype or "trade" in stype:
        if journal:
            parts.append(f"*{journal}.*")
        if doi.startswith("http"):
            parts.append(doi)
    elif journal:
        if vol:
            ref = f"*{journal}, {vol}*"
            if iss:
                ref += f"({iss})"
            if pg:
                ref += f", {pg}"
            parts.append(ref + ".")
        elif s.get("online_first"):
            parts.append(f"*{journal}.* Advance online publication.")
        else:
            parts.append(f"*{journal}.*")
        if doi.startswith("https://doi.org/"):
            parts.append(doi)
        elif doi.startswith("http"):
            parts.append(doi)
        elif s.get("pmid"):
            parts.append(f"https://pubmed.ncbi.nlm.nih.gov/{s['pmid']}/")
    return " ".join(x for x in parts if x)


def sort_key(s):
    first = (s.get("authors") or ["zzz"])
    name = first[0].lower() if first else "zzzz"
    return (name, str(s.get("year") or "0"))


def smoke_test():
    """2 dummy sources -> build -> verify no double periods and non-trivial output."""
    tmp = tempfile.mkdtemp(prefix="apa_smoke_")
    sources = [
        {"authors": ["Doe, J. A.", "Smith, B. C."], "year": 2020, "title": "A study of something important",
         "journal": "Journal of Testing", "volume": "12", "issue": "3", "pages": "45-67",
         "doi_or_url": "https://doi.org/10.1000/fake.doi", "pmid": "", "type": "peer-reviewed",
         "summary": "This is a smoke-test summary long enough to pass the sixty character threshold easily."},
        {"authors": ["Home Office"], "year": None, "title": "Guidance for testing",
         "journal": "GOV.UK", "volume": "", "issue": "", "pages": "",
         "doi_or_url": "https://www.gov.uk/test", "pmid": "", "type": "guideline/report",
         "summary": "Another smoke-test summary that is definitely long enough to pass validation checks."},
    ]
    with open(os.path.join(tmp, "sources.json"), "w", encoding="utf-8") as f:
        json.dump(sources, f)
    with open(os.path.join(tmp, "synthesis.md"), "w", encoding="utf-8") as f:
        f.write("## Introduction\n\nBody text here.\n")
    out = os.path.join(tmp, "smoke.docx")
    build(tmp, out, "Smoke Test Document")
    d = Document(out)
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Doe, J. A., & Smith, B. C. (2020)" in text, "author formatting broken"
    assert ".." not in text.replace("...", ""), "double-period regression"
    assert "Home Office. (n.d.)" in text, "org-author handling broken"
    print(f"SMOKE TEST PASS -> {out}")
    return 0


def build(ws_dir, out_path, title):
    with open(os.path.join(ws_dir, "sources.json"), encoding="utf-8") as f:
        sources = json.load(f)
    seen, uniq = set(), []
    for s in sources:
        k = (s.get("doi_or_url") or s.get("title") or "").lower().strip()
        if k and k in seen:
            continue
        seen.add(k); uniq.append(s)
    sources = sorted(uniq, key=sort_key)

    synth_path = os.path.join(ws_dir, "synthesis.md")
    synth_md = open(synth_path, encoding="utf-8").read() if os.path.exists(synth_path) else ""
    abs_path = os.path.join(ws_dir, "abstract.md")
    abstract = open(abs_path, encoding="utf-8").read().strip() if os.path.exists(abs_path) else ""

    doc = Document()
    set_apa_styles(doc)
    add_page_number_header(doc)
    doc.core_properties.title = title

    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, title, bold=True)
    doc.add_paragraph()
    add_centered(doc, "A Literature Review")
    doc.add_paragraph()
    add_centered(doc, MONTH_YEAR)
    doc.add_page_break()

    add_h1(doc, "Abstract")
    p = doc.add_paragraph()
    _emit_runs(p, abstract or f"This review synthesizes {len(sources)} sources on the topic below.")
    kw = doc.add_paragraph()
    kr = kw.add_run("Keywords: "); kr.italic = True
    kw.add_run("literature review, workplace prevention, sexual assault, domestic violence")
    doc.add_page_break()

    for line in synth_md.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#### "):
            p = doc.add_paragraph(); r = p.add_run(s[5:]); r.bold = True; r.italic = True
        elif s.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(s[4:]); r.bold = True
        elif s.startswith("## "):
            add_h1(doc, s[3:])
        elif s.startswith("# "):
            continue
        else:
            add_body(doc, s)
    doc.add_page_break()

    add_h1(doc, "References")
    for s in sources:
        hanging(doc, ref_entry(s))
    doc.add_page_break()

    add_h1(doc, "Annotated Bibliography")
    intro = doc.add_paragraph()
    _emit_runs(intro, f"The following section provides a brief summary of every source cited in this review ({len(sources)} sources).")
    doc.add_paragraph()
    for s in sources:
        hanging(doc, ref_entry(s))
        summ = doc.add_paragraph()
        summ.paragraph_format.left_indent = Inches(0.5)
        stype = s.get("type", "")
        _emit_runs(summ, f"[{stype}] {(s.get('summary') or '').strip()}")
        doc.add_paragraph()

    doc.save(out_path)
    print(f"Saved: {out_path}  ({len(sources)} sources)")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--smoke-test":
        sys.exit(smoke_test())
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    build(sys.argv[1], sys.argv[2], sys.argv[3])
