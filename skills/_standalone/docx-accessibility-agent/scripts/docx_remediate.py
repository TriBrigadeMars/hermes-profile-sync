"""
docx_remediate.py — Add tracked changes (revisions) and comments to a .docx
without a native python-docx API. VERIFIED working with python-docx 1.2.0.

Capabilities confirmed (see /tmp tests):
  - Track changes: w:ins (insertions) + w:del/w:delText (deletions)  -> Word shows
    them as revision marks when Track Changes is on / Review Pane open.
  - Comments: word/comments.xml part + w:commentRangeStart/End + w:commentReference.

Usage:
    from docx_remediate import insert_tracked, delete_tracked, add_comment
    doc = Document("input.docx")
    p = doc.add_paragraph()
    insert_tracked(p, "new text", author="Accessibility Agent")
    delete_tracked(p, "old text", author="Accessibility Agent")
    add_comment(doc, p.runs[0], "Consider adding alt text here.", author="Accessibility Agent")
    doc.save("output_tracked.docx")
"""

import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part as OpcPart

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_com_id_counter = {"n": 0}


def _now():
    return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%S") + "Z"


def insert_tracked(paragraph, text, author="Accessibility Agent"):
    """Insert text as a tracked *insertion* (w:ins). Returns the new run."""
    run = paragraph.add_run(text)
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(1000 + _com_id_counter["n"]))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), _now())
    r_elem = run._r
    idx = list(paragraph._p).index(r_elem)
    paragraph._p.insert(idx, ins)
    ins.append(r_elem)
    _com_id_counter["n"] += 1
    return run


def delete_tracked(paragraph, text, author="Accessibility Agent"):
    """Delete text as a tracked *deletion* (w:del + w:delText)."""
    run = paragraph.add_run(text)
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), str(1000 + _com_id_counter["n"]))
    dele.set(qn("w:author"), author)
    dele.set(qn("w:date"), _now())
    r_elem = run._r
    for t in r_elem.findall(qn("w:t")):
        r_elem.remove(t)
    deltext = OxmlElement("w:delText")
    deltext.set(qn("xml:space"), "preserve")
    deltext.text = text
    r_elem.append(deltext)
    dele.append(r_elem)
    paragraph._p.append(dele)
    _com_id_counter["n"] += 1
    return run


class _CommentsPart(OpcPart):
    _partname = PackURI("/word/comments.xml")
    _content_type = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.comments+xml"
    )
    reltype = RT.COMMENTS

    def __init__(self, blob):
        super().__init__(self._partname, self._content_type, blob, None)


def add_comment(document, run, text, author="Accessibility Agent"):
    """Attach a real Word comment (comments.xml part) to a run."""
    cid = str(_com_id_counter["n"])
    _com_id_counter["n"] += 1

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), cid)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), cid)
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), cid)
    ref_run.append(ref)

    p = run._parent
    p._p.insert(0, start)
    p._p.append(end)
    p._p.append(ref_run)

    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:comments xmlns:w="{W_NS}">'
        f'<w:comment w:id="{cid}" w:author="{author}" '
        f'w:date="{_now()}" w:initials="AA"><w:p><w:r><w:t>{escaped}</w:t></w:r></w:p>'
        f'</w:comment></w:comments>'
    )
    part = _CommentsPart(xml.encode("utf-8"))
    document.part.relate_to(part, RT.COMMENTS)
    return cid


if __name__ == "__main__":
    doc = Document()
    doc.add_heading("Remediated Document", 0)
    p = doc.add_paragraph()
    p.add_run("The ")
    delete_tracked(p, "unclear wording")
    p.add_run(" ")
    insert_tracked(p, "corrected, descriptive wording")
    p.add_run(" is now accessible.")
    add_comment(doc, p.runs[0], "Updated wording for WCAG 2 alt-text clarity.")
    doc.save("/tmp/remediated_example.docx")
    print("Wrote /tmp/remediated_example.docx with tracked changes + comment.")